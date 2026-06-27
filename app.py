import os
import io
import json
import time
import uuid
import base64
import sqlite3
import re
import logging
import traceback
from datetime import datetime
from functools import wraps
from flask import (
    Flask, render_template, request, jsonify,
    send_file, send_from_directory, Response, stream_with_context, after_this_request
)
from werkzeug.utils import secure_filename
from flask_cors import CORS

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(os.path.dirname(__file__), 'ocr_app.log')),
        logging.StreamHandler(),
    ]
)
log = logging.getLogger(__name__)

from ocr_preprocess import (
    preprocess_for_ocr, postprocess_khmer,
    detect_image_quality, correct_with_ai
)

app = Flask(__name__)
CORS(app, origins=["http://localhost:5173", "http://localhost:5001"])
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50 MB
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')

BASE_DIR         = os.path.dirname(os.path.abspath(__file__))
CORRECTIONS_FILE = os.path.join(BASE_DIR, 'corrections.json')
DB_PATH          = os.path.join(BASE_DIR, 'history.db')

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# ── In-memory result store (capped) ────────────────────────────────
results_store = {}
_MAX_RESULTS  = 100

def _store_result(file_id, result):
    if len(results_store) >= _MAX_RESULTS:
        del results_store[next(iter(results_store))]
    results_store[file_id] = result


# ── .env loader ─────────────────────────────────────────────────────
def _load_env():
    env_path = os.path.join(BASE_DIR, '.env')
    if os.path.exists(env_path):
        with open(env_path) as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    k, v = line.split('=', 1)
                    os.environ.setdefault(k.strip(), v.strip().strip("'\""))

_load_env()


# ── SQLite history ──────────────────────────────────────────────────
def _init_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute('''
        CREATE TABLE IF NOT EXISTS history (
            id           TEXT PRIMARY KEY,
            filename     TEXT,
            timestamp    TEXT,
            preset       TEXT,
            ai_corrected INTEGER,
            total_pages  INTEGER,
            char_count   INTEGER,
            full_text    TEXT,
            pages_json   TEXT
        )
    ''')
    conn.commit()
    conn.close()

_init_db()


def _save_history(result):
    try:
        pages_no_thumb = [
            {k: v for k, v in p.items() if k != 'thumbnail'}
            for p in result.get('pages', [])
        ]
        conn = sqlite3.connect(DB_PATH)
        conn.execute(
            'INSERT OR REPLACE INTO history VALUES (?,?,?,?,?,?,?,?,?)',
            (
                result['file_id'],
                result['filename'],
                datetime.now().isoformat(),
                result.get('preset', 'auto'),
                1 if result.get('ai_corrected') else 0,
                result['total_pages'],
                len(result.get('full_text', '')),
                result.get('full_text', ''),
                json.dumps(pages_no_thumb, ensure_ascii=False),
            )
        )
        conn.commit()
        conn.close()
    except Exception as e:
        print(f'⚠️  History save failed: {e}')


# ── User corrections ────────────────────────────────────────────────
def _load_corrections():
    if os.path.exists(CORRECTIONS_FILE):
        try:
            with open(CORRECTIONS_FILE, encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return []
    return []


def _save_corrections(corrections):
    with open(CORRECTIONS_FILE, 'w', encoding='utf-8') as f:
        json.dump(corrections, f, ensure_ascii=False, indent=2)


def _apply_user_corrections(text):
    for c in _load_corrections():
        orig = c.get('original', '')
        corr = c.get('corrected', '')
        if orig and corr:
            text = text.replace(orig, corr)
    return text


# ── API-key auth decorator ──────────────────────────────────────────
def require_api_key(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        key      = request.headers.get('X-API-Key') or request.args.get('api_key')
        expected = os.environ.get('ADMIN_API_KEY')
        if not expected:
            return jsonify({'error': 'REST API disabled — set ADMIN_API_KEY in .env'}), 503
        if key != expected:
            return jsonify({'error': 'Invalid or missing X-API-Key header'}), 401
        return f(*args, **kwargs)
    return decorated


# ── File-type helpers ───────────────────────────────────────────────
ALLOWED_EXTENSIONS = {'pdf', 'jpg', 'jpeg', 'png', 'tiff', 'tif', 'bmp', 'webp'}
IMAGE_EXTENSIONS   = {'jpg', 'jpeg', 'png', 'tiff', 'tif', 'bmp', 'webp'}

def _ext(filename):
    return filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

def allowed_file(filename):
    return _ext(filename) in ALLOWED_EXTENSIONS

def is_image(filename):
    return _ext(filename) in IMAGE_EXTENSIONS


# ── Image utilities ─────────────────────────────────────────────────
def _page_thumbnail(pil_img, max_w=560):
    from PIL import Image
    thumb = pil_img.copy()
    w, h  = thumb.size
    if w > max_w:
        thumb = thumb.resize((max_w, int(h * max_w / w)), resample=Image.LANCZOS)
    buf = io.BytesIO()
    thumb.convert('RGB').save(buf, format='JPEG', quality=65)
    return 'data:image/jpeg;base64,' + base64.b64encode(buf.getvalue()).decode()


def _confidence(text, word_confidences=None):
    # Use real Tesseract word-level confidence scores when available
    if word_confidences:
        scores = [c for c in word_confidences.values() if c >= 0]
        if scores:
            return round(min(1.0, sum(scores) / len(scores) / 100), 2)
    # Fallback: Khmer character ratio (used when no Tesseract data)
    chars = [c for c in text if c not in (' ', '\n', '\t')]
    if not chars:
        return 0.0
    khmer = sum(1 for c in chars if 'ក' <= c <= '៿')
    return round(min(1.0, khmer / len(chars) * 1.3), 2)


def _sse(data):
    return f"data: {json.dumps(data, ensure_ascii=False)}\n\n"


# ── Core per-page processor ─────────────────────────────────────────
def _process_page(img, page_num, total_pages, preset):
    import pytesseract
    import numpy as np
    import cv2

    tesseract_path = '/opt/homebrew/bin/tesseract'
    if os.path.exists(tesseract_path):
        pytesseract.pytesseract.tesseract_cmd = tesseract_path

    config = '--oem 1 --psm 4 -c preserve_interword_spaces=1 -c textord_min_linesize=1.5'

    # Per-stage timing for performance evaluation
    t_preprocess_start = time.perf_counter()

    log.info(f'[page {page_num}] thumbnail...')
    thumbnail  = _page_thumbnail(img)
    log.info(f'[page {page_num}] to array...')
    gray       = np.array(img)
    if len(gray.shape) == 3:
        gray = cv2.cvtColor(gray, cv2.COLOR_RGB2GRAY)
    log.info(f'[page {page_num}] image size: {gray.shape}, quality detect...')
    quality        = detect_image_quality(gray)
    log.info(f'[page {page_num}] preset={quality["recommended_preset"]}, preprocess...')
    processed      = preprocess_for_ocr(img, preset=preset)
    log.info(f'[page {page_num}] preprocess done, size={np.array(processed).shape}, tesseract...')
    t_preprocess_end = time.perf_counter()
    t_tess_start = t_preprocess_end

    def _run_tess(cfg, timeout):
        from pytesseract import Output as TessOut
        data = pytesseract.image_to_data(
            processed, lang='khm+eng', config=cfg,
            output_type=TessOut.DICT, timeout=timeout
        )
        lines, confs = {}, {}
        for i, word in enumerate(data['text']):
            w = str(word).strip()
            c = int(data['conf'][i])
            if not w:
                continue
            key = (data['block_num'][i], data['par_num'][i], data['line_num'][i])
            lines.setdefault(key, []).append(w)
            if c >= 0:
                wl = w.lower()
                if wl not in confs or c < confs[wl]:
                    confs[wl] = c
        return '\n'.join(' '.join(ws) for ws in lines.values()), confs

    raw_text, word_confidences = '', {}
    try:
        raw_text, word_confidences = _run_tess(config, 25)
    except RuntimeError:
        log.warning(f'[page {page_num}] Tesseract timed out, retrying with psm 3...')
        try:
            raw_text, word_confidences = _run_tess('--oem 1 --psm 3', 20)
        except (RuntimeError, Exception):
            log.warning(f'[page {page_num}] Fallback also timed out, skipping page')
    except Exception:
        pass
    log.info(f'[page {page_num}] tesseract done, chars={len(raw_text)}')
    t_tess_end = time.perf_counter()
    t_post_start = t_tess_end

    corrected_text = postprocess_khmer(raw_text)
    corrected_text = _apply_user_corrections(corrected_text)
    t_post_end = time.perf_counter()

    # Suspicious words = low confidence AND not auto-corrected by postprocess.
    # Pure confidence flags too many correct Khmer words (Tesseract scores Khmer low
    # even when correct). Diffing raw vs corrected lets us exclude words that were
    # already fixed, leaving only genuinely uncertain ones.
    import difflib
    raw_tokens  = [w for w in raw_text.split()       if w.strip()]
    corr_tokens = [w for w in corrected_text.split() if w.strip()]
    unchanged_in_corr = set()
    for tag, i1, i2, j1, j2 in difflib.SequenceMatcher(
        None,
        [w.lower() for w in raw_tokens],
        [w.lower() for w in corr_tokens],
        autojunk=False,
    ).get_opcodes():
        if tag == 'equal':
            unchanged_in_corr.update(w.lower() for w in corr_tokens[j1:j2])
    # Only flag unchanged words whose Tesseract confidence is very low
    suspicious_words = {
        w: c for w, c in word_confidences.items()
        if w in unchanged_in_corr and c < 45
    }

    raw_lines  = len(raw_text.strip().split('\n'))
    corr_lines = len(corrected_text.strip().split('\n'))

    return {
        'page':             page_num,
        'text':             corrected_text,
        'raw_text':         raw_text,
        'quality':          quality,
        'lines_removed':    max(0, raw_lines - corr_lines),
        'char_count':       len(corrected_text),
        'confidence':       _confidence(corrected_text, word_confidences),
        'thumbnail':        thumbnail,
        'word_confidences': suspicious_words,
        'timings': {
            'preprocess_s':  round(t_preprocess_end - t_preprocess_start, 2),
            'tesseract_s':   round(t_tess_end - t_tess_start, 2),
            'postprocess_s': round(t_post_end - t_post_start, 2),
            'total_s':       round(t_post_end - t_preprocess_start, 2),
        },
    }


# ═══════════════════════════════════════════════════════════════════
# ROUTES
# ═══════════════════════════════════════════════════════════════════

@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def index(path):
    # In production, serve the built React app from frontend/dist/
    dist_dir = os.path.join(BASE_DIR, 'frontend', 'dist')
    if os.path.exists(dist_dir):
        target = os.path.join(dist_dir, path)
        if path and os.path.exists(target) and os.path.isfile(target):
            return send_from_directory(dist_dir, path)
        return send_from_directory(dist_dir, 'index.html')
    # Local dev fallback — Flask template
    return render_template('index.html')


# ── Upload ──────────────────────────────────────────────────────────
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    if not allowed_file(file.filename):
        return jsonify({'error': 'Unsupported type. Allowed: PDF, JPG, PNG, TIFF, BMP, WEBP'}), 400

    file_id  = str(uuid.uuid4())[:8]
    filename = secure_filename(file.filename)
    path     = os.path.join(app.config['UPLOAD_FOLDER'], f'{file_id}_{filename}')
    file.save(path)
    return jsonify({'file_id': file_id, 'filename': filename, 'size': os.path.getsize(path)})


# ── Process (SSE stream) ────────────────────────────────────────────
@app.route('/process', methods=['POST'])
def process_file():
    data        = request.json
    file_id     = data.get('file_id')
    preset      = data.get('preset', 'auto')
    use_ai      = data.get('use_ai', False)

    upload_dir = app.config['UPLOAD_FOLDER']
    matching   = [f for f in os.listdir(upload_dir) if f.startswith(file_id)]
    if not matching:
        return jsonify({'error': 'File not found'}), 404

    filepath = os.path.join(upload_dir, matching[0])

    def generate():
        from PIL import Image

        log.info(f'[{file_id}] Starting processing: {matching[0]}, preset={preset}, use_ai={use_ai}')
        yield _sse({'step': 'converting', 'message': '📄 Loading document...'})
        try:
            if is_image(matching[0]):
                images = [Image.open(filepath)]
                total_pages = 1
                log.info(f'[{file_id}] Image file, 1 page')
            else:
                import subprocess
                try:
                    pdfinfo_bin = '/opt/homebrew/bin/pdfinfo'
                    info = subprocess.run(
                        [pdfinfo_bin, filepath], capture_output=True, text=True, timeout=10
                    )
                    pages_line = next(
                        (l for l in info.stdout.splitlines() if l.startswith('Pages:')), None
                    )
                    total_pages = int(pages_line.split(':')[1].strip()) if pages_line else None
                    log.info(f'[{file_id}] pdfinfo found {total_pages} pages')
                except Exception as pe:
                    log.warning(f'[{file_id}] pdfinfo failed ({pe}), will full-convert')
                    total_pages = None
                images = None
        except Exception as e:
            log.error(f'[{file_id}] Load error: {traceback.format_exc()}')
            yield _sse({'step': 'error', 'message': f'❌ Failed to load: {e}'})
            return

        if images is not None:
            all_images = images
            total_pages = 1
        else:
            if total_pages is None:
                from pdf2image import convert_from_path
                log.info(f'[{file_id}] Full-converting PDF (pdfinfo unavailable)...')
                all_images = convert_from_path(filepath, dpi=300)
                total_pages = len(all_images)
                log.info(f'[{file_id}] Full-convert done: {total_pages} pages')
            else:
                all_images = None

        yield _sse({'step': 'converted', 'message': f'📄 Found {total_pages} page(s)', 'total_pages': total_pages})

        pages_data = []
        for pn in range(1, total_pages + 1):
            if all_images is not None:
                img = all_images[pn - 1]
            else:
                from pdf2image import convert_from_path
                log.info(f'[{file_id}] Converting page {pn}/{total_pages}...')
                converted = convert_from_path(filepath, dpi=300, first_page=pn, last_page=pn)
                if not converted:
                    log.error(f'[{file_id}] convert_from_path returned empty for page {pn}')
                    yield _sse({'step': 'error', 'message': f'❌ Failed to convert page {pn}'})
                    return
                img = converted[0]

            log.info(f'[{file_id}] OCR page {pn}/{total_pages}...')
            yield _sse({'step': 'preprocessing', 'message': f'🔧 Page {pn}/{total_pages}: Preprocessing...', 'page': pn})
            try:
                page = _process_page(img, pn, total_pages, preset)
            except Exception as e:
                log.error(f'[{file_id}] Page {pn} failed: {traceback.format_exc()}')
                yield _sse({'step': 'error', 'message': f'❌ Page {pn} error: {e}'})
                return
            pages_data.append(page)
            log.info(f'[{file_id}] Page {pn} done, chars={page["char_count"]}')

            yield _sse({
                'step':     'page_done',
                'message':  f'✅ Page {pn}/{total_pages} complete',
                'page':     pn,
                'progress': round(pn / total_pages * 90),
            })

        full_text = '\n'.join(p['text'] for p in pages_data)

        if use_ai:
            yield _sse({'step': 'ai_correction', 'message': '🤖 Applying AI correction (Google Gemini)...'})
            try:
                full_text = correct_with_ai(full_text)
                yield _sse({'step': 'ai_done', 'message': '🤖 AI correction complete!'})
            except Exception as e:
                yield _sse({'step': 'ai_error', 'message': f'⚠️ AI correction failed: {e}'})

        result = {
            'file_id':      file_id,
            'filename':     matching[0].split('_', 1)[1] if '_' in matching[0] else matching[0],
            'total_pages':  total_pages,
            'pages':        pages_data,
            'full_text':    full_text,
            'preset':       preset,
            'ai_corrected': use_ai,
        }
        _store_result(file_id, result)
        _save_history(result)

        yield _sse({'step': 'complete', 'message': f'🎉 All {total_pages} page(s) processed!', 'result': result})

        try:
            os.remove(filepath)
            print(f'🧹 Auto-deleted: {filepath}')
        except Exception:
            pass

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'}
    )


# ── Download TXT ────────────────────────────────────────────────────
@app.route('/download/<file_id>/txt')
@app.route('/download/<file_id>')          # backwards-compat
def download_txt(file_id):
    if file_id not in results_store:
        return jsonify({'error': 'No results found for this file'}), 404
    result   = results_store[file_id]
    txt_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{file_id}_output.txt')
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(result['full_text'])
    name = result['filename'].rsplit('.', 1)[0] + '_ocr.txt'

    @after_this_request
    def _cleanup(resp):
        try: os.remove(txt_path)
        except Exception: pass
        return resp

    return send_file(txt_path, as_attachment=True, download_name=name)


# ── Download DOCX ───────────────────────────────────────────────────
@app.route('/download/<file_id>/docx')
def download_docx(file_id):
    if file_id not in results_store:
        return jsonify({'error': 'No results found for this file'}), 404
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return jsonify({'error': 'python-docx not installed. Run: pip install python-docx'}), 500

    result = results_store[file_id]
    doc    = Document()

    # Title
    title             = doc.add_heading(result['filename'].rsplit('.', 1)[0], 0)
    title.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    sub               = doc.add_paragraph(
        f"Khmer OCR Extraction  ·  {result['total_pages']} page(s)  ·  "
        f"{datetime.now().strftime('%Y-%m-%d')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for page in result['pages']:
        doc.add_heading(f"Page {page['page']}", level=2)
        for para_text in page['text'].split('\n\n'):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                for run in p.runs:
                    run.font.name = 'Khmer OS'
                    run.font.size = Pt(13)
        doc.add_paragraph()

    docx_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{file_id}_output.docx')
    doc.save(docx_path)
    name = result['filename'].rsplit('.', 1)[0] + '_ocr.docx'

    @after_this_request
    def _cleanup(resp):
        try: os.remove(docx_path)
        except Exception: pass
        return resp

    return send_file(docx_path, as_attachment=True, download_name=name)


# ── Inline AI correction (Gemini or Claude) ────────────────────────
@app.route('/correct', methods=['POST'])
def correct_selection():
    data     = request.json
    selected = (data.get('selected_text') or '').strip()
    if not selected:
        return jsonify({'error': 'No text selected'}), 400

    provider = os.environ.get('AI_PROVIDER', 'gemini').lower()
    prompt = (
        "You are a Khmer OCR correction expert.\n"
        "The text below was extracted from a scanned PDF using Tesseract OCR.\n\n"
        f"Text to analyze:\n{selected}\n\n"
        'Return ONLY a JSON object: {"corrections":[{"original":"exact wrong text",'
        '"corrected":"correct replacement","reason":"brief explanation"}]}\n\n'
        "Rules:\n"
        "- 'original' must be copied exactly from the input\n"
        "- Only real OCR errors (wrong chars, missing diacritics)\n"
        "- No stylistic changes\n"
        "- If text is correct, return {\"corrections\":[]}\n"
        "- Return ONLY the JSON, no markdown, no code fences"
    )

    try:
        if provider == 'claude':
            api_key = os.environ.get('ANTHROPIC_API_KEY')
            if not api_key:
                return jsonify({'error': 'ANTHROPIC_API_KEY not set in .env file'}), 500
            import anthropic as _anthropic
            _client = _anthropic.Anthropic(api_key=api_key)
            response = _client.messages.create(
                model='claude-haiku-4-5',
                max_tokens=2048,
                messages=[{'role': 'user', 'content': prompt}]
            )
            raw = response.content[0].text.strip()
        else:
            api_key = os.environ.get('GOOGLE_API_KEY')
            if not api_key:
                return jsonify({'error': 'GOOGLE_API_KEY not set in .env file'}), 500
            from google import genai as _genai
            _client = _genai.Client(api_key=api_key)
            response = _client.models.generate_content(model='gemini-2.0-flash', contents=prompt)
            raw = response.text.strip()

        if raw.startswith('```'):
            raw = raw.split('\n', 1)[1].rsplit('```', 1)[0].strip()
        return jsonify(json.loads(raw))
    except json.JSONDecodeError:
        return jsonify({'error': 'AI returned invalid JSON. Try again.'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ── Documents listing & local processing ───────────────────────────
@app.route('/documents')
def list_documents():
    doc_dir = os.path.join(BASE_DIR, 'Document')
    if not os.path.exists(doc_dir):
        return jsonify([])
    docs = []
    for f in sorted(os.listdir(doc_dir)):
        if f.lower().endswith('.pdf'):
            path = os.path.join(doc_dir, f)
            docs.append({'name': f, 'size': os.path.getsize(path)})
    return jsonify(docs)


@app.route('/process-local', methods=['POST'])
def process_local_file():
    import shutil
    data     = request.json
    filename = os.path.basename(data.get('filename', ''))
    if not filename or not filename.lower().endswith('.pdf'):
        return jsonify({'error': 'Invalid filename'}), 400
    doc_path = os.path.join(BASE_DIR, 'Document', filename)
    if not os.path.exists(doc_path):
        return jsonify({'error': 'Document not found'}), 404
    file_id  = str(uuid.uuid4())[:8]
    dst      = os.path.join(app.config['UPLOAD_FOLDER'], f'{file_id}_{filename}')
    shutil.copy2(doc_path, dst)
    return jsonify({'file_id': file_id, 'filename': filename, 'size': os.path.getsize(dst)})


# ── History ─────────────────────────────────────────────────────────
@app.route('/history')
def get_history():
    try:
        conn = sqlite3.connect(DB_PATH)
        rows = conn.execute(
            'SELECT id, filename, timestamp, preset, ai_corrected, total_pages, char_count '
            'FROM history ORDER BY timestamp DESC LIMIT 50'
        ).fetchall()
        conn.close()
        return jsonify([{
            'id': r[0], 'filename': r[1], 'timestamp': r[2],
            'preset': r[3], 'ai_corrected': bool(r[4]),
            'total_pages': r[5], 'char_count': r[6]
        } for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/history/<history_id>')
def get_history_item(history_id):
    try:
        conn = sqlite3.connect(DB_PATH)
        row  = conn.execute('SELECT * FROM history WHERE id=?', (history_id,)).fetchone()
        conn.close()
        if not row:
            return jsonify({'error': 'Not found'}), 404
        return jsonify({
            'id': row[0], 'filename': row[1], 'timestamp': row[2],
            'preset': row[3], 'ai_corrected': bool(row[4]),
            'total_pages': row[5], 'char_count': row[6],
            'full_text': row[7],
            'pages': json.loads(row[8]) if row[8] else []
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/history/<history_id>', methods=['DELETE'])
def delete_history_item(history_id):
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.execute('DELETE FROM history WHERE id=?', (history_id,))
        conn.commit()
        conn.close()
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ── User corrections ────────────────────────────────────────────────
@app.route('/corrections', methods=['GET'])
def get_corrections():
    return jsonify(_load_corrections())


@app.route('/corrections', methods=['POST'])
def add_correction():
    data      = request.json
    original  = (data.get('original') or '').strip()
    corrected = (data.get('corrected') or '').strip()
    if not original or not corrected:
        return jsonify({'error': 'Both original and corrected are required'}), 400
    if original == corrected:
        return jsonify({'error': 'original and corrected are the same'}), 400

    corrections = _load_corrections()
    for c in corrections:
        if c['original'] == original:
            c['corrected'] = corrected
            c['updated']   = datetime.now().isoformat()
            break
    else:
        corrections.append({
            'original':  original,
            'corrected': corrected,
            'added':     datetime.now().isoformat()
        })
    _save_corrections(corrections)
    return jsonify({'ok': True, 'total': len(corrections)})


@app.route('/corrections/<int:idx>', methods=['DELETE'])
def delete_correction(idx):
    corrections = _load_corrections()
    if idx < 0 or idx >= len(corrections):
        return jsonify({'error': 'Index out of range'}), 404
    corrections.pop(idx)
    _save_corrections(corrections)
    return jsonify({'ok': True})


# ── REST API (for developers) ────────────────────────────────────────
@app.route('/api/v1/extract', methods=['POST'])
@require_api_key
def api_extract():
    """
    Synchronous REST extraction endpoint.
    Headers  : X-API-Key: <ADMIN_API_KEY>
    Body     : multipart/form-data, field 'file' (PDF or image)
    Optional : preset=auto|clean|scan|photo  use_ai=true|false
    Returns  : JSON result (no thumbnails)
    """
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    file    = request.files['file']
    preset  = request.form.get('preset', 'auto')
    use_ai  = request.form.get('use_ai', 'false').lower() == 'true'
    if not allowed_file(file.filename):
        return jsonify({'error': 'Unsupported file type'}), 400

    file_id  = str(uuid.uuid4())[:8]
    filename = secure_filename(file.filename)
    path     = os.path.join(app.config['UPLOAD_FOLDER'], f'{file_id}_{filename}')
    file.save(path)
    start = time.time()

    try:
        from PIL import Image
        if is_image(filename):
            images = [Image.open(path)]
        else:
            from pdf2image import convert_from_path
            images = convert_from_path(path, dpi=300)

        pages_data = []
        for i, img in enumerate(images):
            page = _process_page(img, i + 1, len(images), preset)
            pages_data.append({k: v for k, v in page.items() if k != 'thumbnail'})

        full_text = '\n'.join(p['text'] for p in pages_data)
        if use_ai:
            full_text = correct_with_ai(full_text)

        result = {
            'file_id':         file_id,
            'filename':        filename,
            'total_pages':     len(images),
            'pages':           pages_data,
            'full_text':       full_text,
            'preset':          preset,
            'ai_corrected':    use_ai,
            'processing_time': round(time.time() - start, 2),
        }
        _store_result(file_id, result)
        _save_history(result)
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try: os.remove(path)
        except Exception: pass


# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    debug = os.environ.get('FLASK_ENV') != 'production'
    print('=' * 50)
    print('  📱 Khmer OCR Web UI')
    print(f'  🌐 Open: http://localhost:{port}')
    print('=' * 50)
    app.run(debug=debug, host='0.0.0.0', port=port, threaded=True)
